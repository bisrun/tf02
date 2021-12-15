from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from konlpy.tag import Hannanum

directory_path= '/home/flywhale/Documents/daimler'
directory_path= 'D:/Documents/++ST++/00_PROOM2020/OEM/다임러/51.SSPICE/프로젝트계획'
docx_filename='Project Management Plan_1.docx'
docx_filepath = directory_path + '/'+docx_filename
hannanum = Hannanum()

"""C:/Anaconda3/envs/tf01/Lib/site-packages/konlpy/java/data/kE/dic_user.txt"
"""

def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


document = Document(docx_filepath)
for block in iter_block_items(document):
    #print(block.text if isinstance(block, Paragraph) else '<table>')
    if isinstance(block, Paragraph):
        print(block.style.name)
        print(block.text)

        if block.text.find("SW요구사양서-리뷰 체크시트") >=0 :
            print("ok")

        #hannanum.morphs(block.text)

    elif isinstance(block, Table):
        print("box"+ block.style.name)
        data = []
        try:
            for row in block.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        data.append(paragraph.text)
                    print("*",data)
                    data = []

        except Exception as e:
            for tc in block._tbl.iter_tcs():
                cell = _Cell(tc, block)
                if tc.left == 0 and len(data) > 0 :
                    print("+", data)
                    data = []

                for b_tc in iter_block_items(cell):
                    if isinstance(b_tc, Paragraph):
                        data.append(b_tc.text)
                    else:
                        data.append(type(b_tc))

            if len(data) > 0 :
                print("+", data)
                data = []

#word 편집제한이 풀려야 함
for paragraph in document.paragraphs:
        hyperlink = paragraph._p.xpath("./w:hyperlink")
        if len(hyperlink) > 0:
            hyperlink = hyperlink[0]
            hyperlink_rel_id = hyperlink.get(qn("r:id"))
            print("hrefid",hyperlink_rel_id)



rels = document.part.rels
def iter_hyperlink_rels(rels):
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            #yield rels[rel]._target
            print('+',rel, rels[rel].target_ref)
        else :
            print('-',rel, rels[rel].target_ref)

print(iter_hyperlink_rels(rels))

print("--end--")
